import io
import logging
import os
import platform
import subprocess
import tempfile
import zipfile

import pdfplumber
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)

# File types we accept (checked in router and here)
SUPPORTED_EXTENSIONS = {"pdf", "docx", "doc", "txt", "rtf"}

_IS_MACOS = platform.system() == "Darwin"


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _extract_pdf(filename, file_bytes)
    elif ext == "docx":
        return _extract_docx(filename, file_bytes)
    elif ext == "doc":
        return _extract_doc(filename, file_bytes)
    elif ext == "rtf":
        return _extract_rtf(filename, file_bytes)
    elif ext == "txt":
        return _extract_txt(filename, file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type '.{ext}' in '{filename}'. "
            f"Accepted formats: PDF, DOCX, DOC, RTF, TXT."
        )


# ── PDF ───────────────────────────────────────────────────────────────────────

def _extract_pdf(filename: str, file_bytes: bytes) -> str:
    """Try three engines in order; raise only when all fail."""

    # Stage 1 — pdfplumber (handles most standard PDFs)
    text = _pdf_pdfplumber(filename, file_bytes)
    if text:
        logger.debug("pdfplumber OK for '%s' (%d chars)", filename, len(text))
        return text

    # Stage 2 — PyMuPDF / fitz (handles more font encodings + compressed streams)
    text = _pdf_pymupdf(filename, file_bytes)
    if text:
        logger.debug("PyMuPDF OK for '%s' (%d chars)", filename, len(text))
        return text

    # Stage 3 — macOS textutil (handles PDFs that embed full CoreText layout)
    if _IS_MACOS:
        text = _textutil_extract(filename, file_bytes, ".pdf")
        if text:
            logger.debug("textutil OK for '%s' (%d chars)", filename, len(text))
            return text

    raise ValueError(
        f"No text could be extracted from '{filename}'. "
        "It may be a scanned/image-only PDF. "
        "Please use a PDF with selectable text, or re-upload the original DOCX/DOC file."
    )


def _pdf_pdfplumber(filename: str, file_bytes: bytes) -> str:
    parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if not text:
                    text = page.extract_text(x_tolerance=10, y_tolerance=10)
                if not text:
                    words = page.extract_words()
                    if words:
                        text = " ".join(w["text"] for w in words)
                if text and text.strip():
                    parts.append(text.strip())
    except Exception as exc:
        err = str(exc).lower()
        if "password" in err or "encrypt" in err:
            raise ValueError(
                f"'{filename}' is password-protected. Please upload an unlocked PDF."
            ) from exc
        logger.warning("pdfplumber failed on '%s': %s", filename, exc)
        return ""
    return "\n".join(parts).strip()


def _pdf_pymupdf(filename: str, file_bytes: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if doc.is_encrypted:
            raise ValueError(
                f"'{filename}' is password-protected. Please upload an unlocked PDF."
            )
        parts = []
        for page in doc:
            text = page.get_text("text")
            if text and text.strip():
                parts.append(text.strip())
        doc.close()
        return "\n".join(parts).strip()
    except ValueError:
        raise
    except Exception as exc:
        logger.warning("PyMuPDF failed on '%s': %s", filename, exc)
        return ""


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _extract_docx(filename: str, file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
    except (PackageNotFoundError, zipfile.BadZipFile) as exc:
        # Possibly a legacy .doc renamed to .docx — try doc path
        logger.warning("'%s' failed as DOCX (%s), retrying as DOC", filename, exc)
        return _extract_doc(filename, file_bytes)
    except Exception as exc:
        logger.exception("python-docx error on '%s'", filename)
        raise ValueError(f"Could not read '{filename}': {exc}.") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in parts:
                    parts.append(t)

    result = "\n".join(parts).strip()
    if not result:
        raise ValueError(f"No text found in '{filename}'. The document appears to be empty.")
    return result


# ── DOC (legacy Word 97-2003) ─────────────────────────────────────────────────

def _extract_doc(filename: str, file_bytes: bytes) -> str:
    """Extract text from .doc files using textutil (macOS) or LibreOffice."""
    if _IS_MACOS:
        text = _textutil_extract(filename, file_bytes, ".doc")
        if text:
            return text

    # Try LibreOffice on Linux/Windows
    text = _libreoffice_extract(filename, file_bytes, ".doc")
    if text:
        return text

    raise ValueError(
        f"Could not convert '{filename}' (legacy Word .doc format). "
        "Please open it in Word/LibreOffice and save as .docx or PDF, then re-upload."
    )


# ── RTF ───────────────────────────────────────────────────────────────────────

def _extract_rtf(filename: str, file_bytes: bytes) -> str:
    # Try macOS textutil first (handles all RTF variants)
    if _IS_MACOS:
        text = _textutil_extract(filename, file_bytes, ".rtf")
        if text:
            return text

    # Fallback: pure-Python striprtf
    try:
        from striprtf.striprtf import rtf_to_text
        raw = file_bytes.decode("utf-8", errors="replace")
        text = rtf_to_text(raw).strip()
        if text:
            return text
    except Exception as exc:
        logger.warning("striprtf failed on '%s': %s", filename, exc)

    raise ValueError(
        f"Could not extract text from RTF file '{filename}'. "
        "Please save it as DOCX or PDF and re-upload."
    )


# ── TXT ───────────────────────────────────────────────────────────────────────

def _extract_txt(filename: str, file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
        try:
            text = file_bytes.decode(encoding).strip()
            if text:
                return text
        except (UnicodeDecodeError, ValueError):
            continue
    raise ValueError(
        f"Could not decode '{filename}'. Please ensure it is a valid plain-text file."
    )


# ── Helpers ───────────────────────────────────────────────────────────────────

def _textutil_extract(filename: str, file_bytes: bytes, suffix: str) -> str:
    """Use macOS textutil to convert a file to plain text."""
    fd, tmp_in = tempfile.mkstemp(suffix=suffix)
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(file_bytes)
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", tmp_in],
            capture_output=True,
            timeout=30,
        )
        if result.returncode == 0:
            text = result.stdout.decode("utf-8", errors="replace").strip()
            if text:
                logger.debug("textutil extracted %d chars from '%s'", len(text), filename)
                return text
        else:
            logger.warning("textutil failed on '%s': %s", filename, result.stderr.decode())
    except subprocess.TimeoutExpired:
        logger.warning("textutil timed out on '%s'", filename)
    except FileNotFoundError:
        logger.debug("textutil not found")
    finally:
        try:
            os.unlink(tmp_in)
        except OSError:
            pass
    return ""


def _libreoffice_extract(filename: str, file_bytes: bytes, suffix: str) -> str:
    """Use LibreOffice headless to convert a file to plain text."""
    fd, tmp_in = tempfile.mkstemp(suffix=suffix)
    out_dir = tempfile.mkdtemp()
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(file_bytes)
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "txt:Text",
             "--outdir", out_dir, tmp_in],
            capture_output=True,
            timeout=60,
        )
        if result.returncode == 0:
            txt_name = os.path.splitext(os.path.basename(tmp_in))[0] + ".txt"
            txt_path = os.path.join(out_dir, txt_name)
            if os.path.exists(txt_path):
                with open(txt_path, "r", errors="replace") as f:
                    text = f.read().strip()
                if text:
                    return text
    except (subprocess.TimeoutExpired, FileNotFoundError):
        logger.debug("libreoffice not available or timed out for '%s'", filename)
    finally:
        try:
            os.unlink(tmp_in)
        except OSError:
            pass
        import shutil
        shutil.rmtree(out_dir, ignore_errors=True)
    return ""
